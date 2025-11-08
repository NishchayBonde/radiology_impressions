import streamlit as st
import openai
import os
import json
import yaml
from datetime import datetime
from pathlib import Path

# Document loaders
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    st.error("python-docx not installed. Install with: pip install python-docx")

try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False

st.set_page_config(page_title="Radiology Impression Generator with Document Loader", layout="wide")

# Configuration files
CONFIG_DIR = Path("config")
CONFIG_DIR.mkdir(exist_ok=True)
STUDIES_CONFIG_FILE = CONFIG_DIR / "studies.yaml"
PROMPTS_FILE = CONFIG_DIR / "study_prompts.json"
DOCUMENTS_DIR = CONFIG_DIR / "study_documents"
DOCUMENTS_DIR.mkdir(exist_ok=True)

# Default studies configuration
DEFAULT_STUDIES = {
    "studies": [
        "CT Chest",
        "CT Abdomen/Pelvis",
        "CT Head/Brain",
        "CT Spine",
        "MRI Brain",
        "MRI Spine",
        "MRI Abdomen",
        "MRI Pelvis",
        "MRI Musculoskeletal",
        "MRI Cardiac",
        "Ultrasound Abdomen",
        "Ultrasound Pelvis",
        "Ultrasound Renal",
        "Ultrasound Cardiac",
        "Chest X-ray",
        "MRCP",
        "CTA",
        "MRA"
    ]
}

# Default prompt template with document placeholder
DEFAULT_PROMPT = """You are an expert radiologist creating a concise radiology impression section following RCR 2018 standards.

CRITICAL INSTRUCTIONS:
1. You MUST follow ALL steps, protocols, and guidelines mentioned in the REFERENCE DOCUMENT below
2. Read the ENTIRE reference document carefully before generating the impression
3. Apply every relevant checklist, measurement standard, and reporting requirement from the document
4. If the document specifies a structure or format, follow it exactly
5. If the document mentions specific terminology or classifications, use them

Generate the impression as exactly 3-4 bullet points:
- Summarize key findings following document protocols
- State primary diagnosis or differential as per document guidelines
- Include any measurements or specifics required by the document

Use clear, standardized terminology as specified in the reference document.

---REFERENCE DOCUMENT START---
{document_content}
---REFERENCE DOCUMENT END---

Now generate the impression following ALL guidelines above."""

# Helper functions
@st.cache_data
def load_studies_from_yaml():
    """Load studies list from YAML file"""
    if STUDIES_CONFIG_FILE.exists():
        try:
            with open(STUDIES_CONFIG_FILE, 'r') as f:
                config = yaml.safe_load(f)
                return config.get('studies', DEFAULT_STUDIES['studies'])
        except Exception as e:
            st.error(f"Error loading studies config: {e}")
            return DEFAULT_STUDIES['studies']
    else:
        with open(STUDIES_CONFIG_FILE, 'w') as f:
            yaml.dump(DEFAULT_STUDIES, f)
        return DEFAULT_STUDIES['studies']

def load_study_prompts():
    """Load study-specific prompts from JSON file"""
    if PROMPTS_FILE.exists():
        try:
            with open(PROMPTS_FILE, 'r') as f:
                return json.load(f)
        except Exception as e:
            st.error(f"Error loading prompts: {e}")
            return {}
    return {}

def save_study_prompt(study_name, prompt):
    """Save prompt for a specific study"""
    prompts = load_study_prompts()
    prompts[study_name] = prompt
    try:
        with open(PROMPTS_FILE, 'w') as f:
            json.dump(prompts, f, indent=2)
        return True
    except Exception as e:
        st.error(f"Error saving prompt: {e}")
        return False

def get_study_prompt(study_name):
    """Get prompt for a specific study"""
    prompts = load_study_prompts()
    return prompts.get(study_name, DEFAULT_PROMPT)

def get_study_document_path(study_name):
    """Get the path for a study's document"""
    safe_name = study_name.replace('/', '_').replace(' ', '_')
    return DOCUMENTS_DIR / f"{safe_name}.docx"

def load_document_content(file_path):
    """Load complete document content with proper structure preservation"""
    if not file_path.exists():
        return ""
    
    try:
        # Method 1: Try docx2txt first (better text extraction)
        if DOCX2TXT_AVAILABLE:
            content = docx2txt.process(str(file_path))
            if content and content.strip():
                return content.strip()
        
        # Method 2: Fallback to python-docx with enhanced structure
        if DOCX_AVAILABLE:
            doc = DocxDocument(file_path)
            
            structured_content = []
            structured_content.append("=" * 80)
            structured_content.append("REFERENCE DOCUMENT - COMPLETE CONTENT")
            structured_content.append("=" * 80)
            structured_content.append("")
            
            current_section = None
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    structured_content.append("")
                    continue
                
                # Detect headings based on style or formatting
                is_heading = False
                if para.style.name.startswith('Heading'):
                    is_heading = True
                    level = para.style.name.replace('Heading', '').strip()
                    if level.isdigit():
                        structured_content.append("")
                        structured_content.append("#" * int(level) + " " + text)
                        structured_content.append("-" * len(text))
                    else:
                        structured_content.append("")
                        structured_content.append("## " + text)
                        structured_content.append("-" * len(text))
                    current_section = text
                # Detect potential headings by characteristics
                elif len(text) < 100 and (text.isupper() or text.endswith(':') or para.runs and para.runs[0].bold):
                    structured_content.append("")
                    structured_content.append("### " + text)
                    structured_content.append("")
                else:
                    structured_content.append(text)
            
            # Add tables if present
            if doc.tables:
                structured_content.append("")
                structured_content.append("=" * 80)
                structured_content.append("TABLES FROM DOCUMENT")
                structured_content.append("=" * 80)
                
                for i, table in enumerate(doc.tables):
                    structured_content.append(f"\n[Table {i+1}]")
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        if row_text:
                            structured_content.append(row_text)
                    structured_content.append("")
            
            return "\n".join(structured_content)
        
        return ""
        
    except Exception as e:
        st.error(f"Error loading document: {e}")
        return ""

def save_study_document(study_name, uploaded_file):
    """Save uploaded document for a specific study"""
    doc_path = get_study_document_path(study_name)
    try:
        with open(doc_path, 'wb') as f:
            f.write(uploaded_file.getbuffer())
        return True
    except Exception as e:
        st.error(f"Error saving document: {e}")
        return False

def count_tokens_estimate(text):
    """Estimate token count (rough approximation)"""
    return len(text.split()) * 1.3  # Approximate token count

# Load studies from YAML
STUDIES_LIST = load_studies_from_yaml()

# Sidebar Configuration
with st.sidebar:
    st.header("⚙️ Configuration")
    openai_api_key = st.text_input("OpenAI API Key", type="password", value=os.getenv("OPENAI_API_KEY", ""))
    
    # Model selection with context window info
    model_options = {
        "gpt-4o (128k context)": "gpt-4o",
        "gpt-4o-mini (128k context)": "gpt-4o-mini",
        "gpt-4-turbo (128k context)": "gpt-4-turbo",
    }
    selected_model = st.selectbox("OpenAI Model", list(model_options.keys()), index=1)
    model = model_options[selected_model]
    
    temperature = st.slider("Temperature", 0.0, 1.0, 0.1, help="Low for consistent reporting")
    max_tokens = st.slider("Max Response Tokens", 500, 2000, 800, help="Maximum length of generated impression")
    
    st.divider()
    
    # Study-specific configuration section
    st.header("📚 Study Configuration")
    config_mode = st.checkbox("Enable Configuration Mode", help="Manage study-specific prompts and documents")
    
    if config_mode:
        st.info("📝 Configure prompts and documents for each study type")
        config_study = st.selectbox("Select Study to Configure", STUDIES_LIST, key="config_study")
        
        with st.expander("📄 Manage Prompt Template", expanded=True):
            st.markdown("""
            **Prompt Variables:**
            - `{document_content}` - Will be replaced with full document text
            - Ensure your prompt instructs the AI to follow ALL document guidelines
            """)
            
            current_prompt = get_study_prompt(config_study)
            edited_prompt = st.text_area(
                f"Prompt Template for {config_study}",
                value=current_prompt,
                height=300,
                key="prompt_editor",
                help="Use {document_content} placeholder where document should be inserted"
            )
            
            # Show token estimate
            if "{document_content}" in edited_prompt:
                doc_path = get_study_document_path(config_study)
                if doc_path.exists():
                    doc_content = load_document_content(doc_path)
                    filled_prompt = edited_prompt.replace("{document_content}", doc_content)
                    est_tokens = count_tokens_estimate(filled_prompt)
                    
                    if est_tokens > 100000:
                        st.error(f"⚠️ Estimated {int(est_tokens):,} tokens - May exceed context window!")
                    elif est_tokens > 50000:
                        st.warning(f"⚠️ Estimated {int(est_tokens):,} tokens - Large context")
                    else:
                        st.success(f"✓ Estimated {int(est_tokens):,} tokens - Within limits")
            
            if st.button("💾 Save Prompt", use_container_width=True):
                if "{document_content}" not in edited_prompt:
                    st.warning("Note: Prompt doesn't contain {document_content} placeholder")
                if save_study_prompt(config_study, edited_prompt):
                    st.success(f"Prompt saved for {config_study}!")
                    st.rerun()
        
        with st.expander("📁 Manage Reference Document", expanded=True):
            doc_path = get_study_document_path(config_study)
            doc_exists = doc_path.exists()
            
            if doc_exists:
                st.success(f"✓ Document exists for {config_study}")
                
                # Show document preview
                doc_content = load_document_content(doc_path)
                doc_tokens = count_tokens_estimate(doc_content)
                
                st.metric("Document Size", f"{len(doc_content):,} chars")
                st.metric("Estimated Tokens", f"{int(doc_tokens):,}")
                
                with st.expander("Preview Document Content"):
                    st.text_area("Content Preview", doc_content[:2000] + "\n\n... (truncated)", height=200, disabled=True)
                
                if st.button("🗑️ Remove Document", use_container_width=True):
                    try:
                        doc_path.unlink()
                        st.success("Document removed!")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Error: {e}")
            else:
                st.warning(f"No document uploaded for {config_study}")
            
            upload_doc = st.file_uploader(
                f"Upload DOCX for {config_study}",
                type="docx",
                key="config_uploader",
                help="Upload complete protocol/guideline document"
            )
            
            if upload_doc and st.button("📤 Upload Document", use_container_width=True):
                if save_study_document(config_study, upload_doc):
                    st.success(f"Document saved for {config_study}!")
                    st.rerun()
        
        st.divider()
        st.caption("💡 Edit studies.yaml to add/remove study types")

st.title("🏥 Radiology Impression Generator with Document Loader")

st.markdown("""
Upload complete protocol documents for each study type. The AI will receive the **entire document** 
and follow **all steps and guidelines** mentioned in it. Prompts are fully customizable from the UI.
""")

col1, col2 = st.columns([2, 3], gap="large")

with col1:
    with st.form("study_form"):
        st.subheader("Study Information")
        
        study_type = st.selectbox("Study Type", STUDIES_LIST)
        
        # Show if study has custom prompt/document
        study_prompt_exists = get_study_prompt(study_type) != DEFAULT_PROMPT
        study_doc_exists = get_study_document_path(study_type).exists()
        
        status_cols = st.columns(2)
        with status_cols[0]:
            if study_prompt_exists:
                st.caption("✓ Custom prompt")
            else:
                st.caption("⚪ Default prompt")
        with status_cols[1]:
            if study_doc_exists:
                st.caption("✓ Reference doc")
            else:
                st.caption("⚪ No document")
        
        # Show document info if available
        if study_doc_exists:
            doc_content = load_document_content(get_study_document_path(study_type))
            doc_tokens = count_tokens_estimate(doc_content)
            st.info(f"📄 Document loaded: ~{int(doc_tokens):,} tokens")
        
        history = st.text_area(
            "Clinical History & Indication", 
            height=80,
            placeholder="Enter clinical history and question..."
        )
        
        findings = st.text_area(
            "Radiologist's Report/Findings", 
            height=150,
            placeholder="Paste key findings..."
        )
        
        submit = st.form_submit_button("Generate Impression", use_container_width=True, type="primary")

# Initialize session state
if 'impression_text' not in st.session_state:
    st.session_state.impression_text = ""
if 'document_used' not in st.session_state:
    st.session_state.document_used = ""
if 'prompt_used' not in st.session_state:
    st.session_state.prompt_used = ""

# Impression generation
if submit and openai_api_key:
    # Load study-specific prompt and document
    study_prompt_template = get_study_prompt(study_type)
    doc_path = get_study_document_path(study_type)
    
    # Load complete document content
    if doc_path.exists():
        with st.spinner("Loading complete reference document..."):
            document_content = load_document_content(doc_path)
            st.session_state.document_used = document_content
    else:
        document_content = "No reference document available for this study type."
        st.session_state.document_used = ""
        st.warning(f"⚠️ No reference document for {study_type}. Using prompt only.")
    
    # Fill in the document content in the prompt template
    system_prompt = study_prompt_template.replace("{document_content}", document_content)
    st.session_state.prompt_used = system_prompt
    
    # Estimate tokens
    total_tokens = count_tokens_estimate(system_prompt)
    
    if total_tokens > 120000:
        st.error(f"❌ Document too large (~{int(total_tokens):,} tokens). Consider using a smaller document or gpt-4o with 128k context.")
    else:
        user_prompt = f"""
Study Type: {study_type}
Clinical History: {history if history else 'Not provided'}
Key Findings: {findings}

Following ALL protocols and guidelines in the reference document above, generate a concise impression in 3-4 bullet points."""
        
        try:
            with st.spinner("Generating impression (processing full document)..."):
                client = openai.OpenAI(api_key=openai_api_key)
                response = client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    temperature=temperature,
                    max_tokens=max_tokens,
                    top_p=0.9
                )
                result = response.choices[0].message.content.strip()
                
                # Add token usage info
                usage = response.usage
                st.session_state.impression_text = result
                st.session_state.token_usage = {
                    "prompt_tokens": usage.prompt_tokens,
                    "completion_tokens": usage.completion_tokens,
                    "total_tokens": usage.total_tokens
                }
            
            st.success("✅ Impression generated!")
            
            # Show actual token usage
            if hasattr(st.session_state, 'token_usage'):
                usage = st.session_state.token_usage
                st.info(f"📊 Tokens used: {usage['prompt_tokens']:,} prompt + {usage['completion_tokens']:,} completion = {usage['total_tokens']:,} total")
            
        except Exception as e:
            st.error(f"Failed to generate: {e}")
            if "context_length_exceeded" in str(e):
                st.error("💡 Try: 1) Use smaller document, 2) Switch to gpt-4o, or 3) Summarize the document first")

# Display
with col2:
    st.subheader("Generated Impression")
    
    impression_area = st.text_area(
        "Editable Impression", 
        st.session_state.impression_text or "Generated impression will appear here...",
        key="edit_impression", 
        height=300
    )
    
    if impression_area != st.session_state.impression_text:
        st.session_state.impression_text = impression_area
    
    # Tabs for different views
    tab1, tab2, tab3 = st.tabs(["📄 Document Used", "📝 Prompt Used", "💾 Export"])
    
    with tab1:
        if st.session_state.document_used:
            doc_preview = st.session_state.document_used
            st.text_area(
                "Complete Reference Document", 
                doc_preview, 
                height=300, 
                disabled=True,
                help="This is the full document content passed to the AI"
            )
            st.download_button(
                "Download Document Content",
                st.session_state.document_used,
                file_name=f"reference_doc_{study_type.replace('/', '_')}.txt"
            )
        else:
            st.info("No document was used for this generation")
    
    with tab2:
        if st.session_state.prompt_used:
            st.text_area(
                "System Prompt Sent to AI", 
                st.session_state.prompt_used[:5000] + "\n\n... (truncated, full prompt was sent)", 
                height=300, 
                disabled=True,
                help="This is the complete prompt including document content"
            )
        else:
            st.info("No prompt available yet")
    
    with tab3:
        col_export1, col_export2 = st.columns(2)
        with col_export1:
            if st.download_button(
                "📥 Download Impression", 
                st.session_state.impression_text, 
                file_name=f"impression_{study_type.replace('/', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                disabled=not st.session_state.impression_text,
                use_container_width=True
            ):
                st.balloons()
        
        with col_export2:
            if st.button("📋 Copy to Clipboard", disabled=not st.session_state.impression_text, use_container_width=True):
                st.code(st.session_state.impression_text)
                st.success("Ready to copy!")

# Footer
st.divider()
st.caption(f"📊 {len(STUDIES_LIST)} study types loaded | 💡 Each study can have custom prompt + full document | ⚙️ Enable Configuration Mode to customize")